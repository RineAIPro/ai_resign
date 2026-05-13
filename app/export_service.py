# 导出服务
# 修改时间：2026/05/12
# 功能：生成交接报告（Word/Markdown/HTML）

import os
import html
from datetime import datetime

from app.database import get_conn


class ExportService:
    def __init__(self, company_id: int, config: dict):
        self.company_id = company_id
        self.config = config

    def generate_report(self, checks: dict, output_path: str) -> str:
        """生成交接报告"""
        ext = os.path.splitext(output_path)[1].lower()

        content = self._build_markdown(checks)

        if ext == ".md":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
        elif ext == ".docx":
            self._save_docx(content, output_path)
        else:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)

        return output_path

    def _build_markdown(self, checks: dict) -> str:
        """构建Markdown内容"""
        # 修改时间：2026/05/12 - 去掉Git提交记录模块
        conn = get_conn()
        cid = self.company_id

        lines = [f"# 离职交接报告", f"", f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", ""]

        # 公司信息
        if checks.get("company"):
            company = conn.execute("SELECT * FROM companies WHERE id=?", (cid,)).fetchone()
            if company:
                lines.extend(["## 公司基本信息", ""])
                lines.append(f"- 公司名称：{company['name']}")
                lines.append(f"- 部门：{company['department']}")
                lines.append(f"- 职位：{company['position']}")
                lines.append(f"- 直属上级：{company['superior']}")
                lines.append(f"- 在职时间：{company['start_date']} ~ {company['leave_date'] or '至今'}")
                lines.append("")

        # 项目清单
        # 修改时间：2026/05/12 - 改为递归遍历，支持无限层级模块
        if checks.get("projects"):
            lines.extend(["## 项目清单", ""])
            modules = conn.execute(
                "SELECT * FROM modules WHERE company_id=? AND parent_id IS NULL", (cid,)
            ).fetchall()
            for mod in modules:
                self._append_module(conn, mod['id'], mod['name'], lines, level=3)

        # 账号
        if checks.get("accounts"):
            lines.extend(["## 账号交接清单", ""])
            accounts = conn.execute("SELECT * FROM accounts WHERE company_id=?", (cid,)).fetchall()
            lines.append("| 系统/平台 | 类型 | 用途 | 状态 |")
            lines.append("| --- | --- | --- | --- |")
            for a in accounts:
                lines.append(f"| {a['platform']} | {a['account_type']} | {a['usage_desc']} | {a['status']} |")
            lines.append("")
            lines.append("> ⚠️ 密码请通过安全渠道单独交接，本报告不含密码明文")
            lines.append("")

        # 待办
        if checks.get("todos"):
            lines.extend(["## 待办事项", ""])
            todos = conn.execute("SELECT * FROM todos WHERE company_id=?", (cid,)).fetchall()
            for t in todos:
                marker = "✅" if t['status'] == '已完成' else "⬜"
                lines.append(f"- {marker} **{t['title']}** — {t['description']}（{t['status']}，{t['priority']}）")
            lines.append("")

        # 联系人
        if checks.get("contacts"):
            lines.extend(["## 联系人", ""])
            contacts = conn.execute("SELECT * FROM contacts WHERE company_id=?", (cid,)).fetchall()
            lines.append("| 姓名 | 角色 | 负责接手 | 联系方式 |")
            lines.append("| --- | --- | --- | --- |")
            for c in contacts:
                lines.append(f"| {c['name']} | {c['role']} | {c['handover_scope']} | {c['contact_info']} |")
            lines.append("")

        # 日程
        if checks.get("schedule"):
            lines.extend(["## 交接日程", ""])
            schedules = conn.execute(
                "SELECT * FROM schedules WHERE company_id=? ORDER BY event_date", (cid,)
            ).fetchall()
            for s in schedules:
                marker = "✅" if s['status'] == '已完成' else "⬜"
                lines.append(f"- {marker} **{s['event_date']}** — {s['content']}（{s['status']}）")
            lines.append("")

        conn.close()
        return "\n".join(lines)

    # 修改时间：2026/05/12 - 项目名改为标题格式
    def _append_projects(self, conn, module_id, lines, level=4):
        projects = conn.execute("SELECT * FROM projects WHERE module_id=?", (module_id,)).fetchall()
        prefix = "#" * level
        for p in projects:
            # 项目名作为标题
            title = p['name']
            status = p['status'] or '未填写'
            if p['version']:
                title += f" v{p['version']}"
            lines.append(f"{prefix} {title}")
            lines.append("")
            lines.append(f"- 状态：{status}")
            tech = p['tech_stack'] or '未填写'
            git = p['git_url'] or '未填写'
            lines.append(f"- 技术栈：{tech}")
            lines.append(f"- Git地址：{git}")
            desc = p['description'] or '未填写'
            lines.append(f"- 简介：{desc}")
            lines.append("")

    # 修改时间：2026/05/12 - 递归遍历模块，支持无限层级
    def _append_module(self, conn, module_id, module_name, lines, level=3):
        prefix = "#" * level
        lines.append(f"{prefix} 📂 {module_name}")
        lines.append("")
        # 当前模块下的项目
        self._append_projects(conn, module_id, lines, level=level + 1)
        # 递归遍历子模块
        subs = conn.execute("SELECT * FROM modules WHERE parent_id=?", (module_id,)).fetchall()
        for sub in subs:
            self._append_module(conn, sub['id'], sub['name'], lines, level=level + 1)

    def _save_docx(self, content: str, output_path: str):
        """保存为Word文档"""
        try:
            from docx import Document
            doc = Document()
            for line in content.split("\n"):
                # 修改时间：2026/05/12 - 改为通用标题解析，支持任意层级
                if line.startswith("#"):
                    # 计算标题层级
                    h_level = 0
                    for ch in line:
                        if ch == '#':
                            h_level += 1
                        else:
                            break
                    if 1 <= h_level <= 9 and len(line) > h_level and line[h_level] == ' ':
                        doc.add_heading(line[h_level + 1:], level=min(h_level, 9))
                elif line.startswith("| ") and "|" in line[1:]:
                    # 表格行，简单处理
                    doc.add_paragraph(line)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith("> "):
                    doc.add_paragraph(line[2:])
                elif line.strip():
                    doc.add_paragraph(line)
            doc.save(output_path)
        except ImportError:
            # 如果没有python-docx，回退保存为md
            md_path = output_path.replace(".docx", ".md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(content)

    # 修改时间：2026/05/12 - 新增项目管理HTML导出
    def generate_projects_html(self) -> str:
        """生成项目清单的独立静态HTML，返回HTML字符串"""
        conn = get_conn()
        cid = self.company_id
        now = datetime.now().strftime('%Y-%m-%d %H:%M')

        company = conn.execute("SELECT * FROM companies WHERE id=?", (cid,)).fetchone()
        company_name = company['name'] if company else '未知'
        department = company['department'] if company else ''

        # 递归统计模块下所有项目数（含子模块）
        def count_descendant_projects(conn, module_id):
            count = conn.execute("SELECT COUNT(*) as c FROM projects WHERE module_id=?", (module_id,)).fetchone()['c']
            subs = conn.execute("SELECT id FROM modules WHERE parent_id=?", (module_id,)).fetchall()
            for sub in subs:
                count += count_descendant_projects(conn, sub['id'])
            return count

        # 递归构建模块HTML
        def build_module_html(module_id, depth=0):
            mod = conn.execute("SELECT * FROM modules WHERE id=?", (module_id,)).fetchone()
            if not mod:
                return ''
            indent = depth * 20
            safe_name = html.escape(mod['name'])

            # 获取项目
            projects = conn.execute("SELECT * FROM projects WHERE module_id=?", (module_id,)).fetchall()
            proj_cards = ''
            for p in projects:
                status = p['status'] or '未填写'
                status_class = 'tag-green' if status == '已完成' else 'tag-orange' if status == '进行中' else 'tag-gray'
                version = f" v{html.escape(p['version'])}" if p['version'] else ''
                tech = html.escape(p['tech_stack'] or '未填写')
                git_url = p['git_url'] or ''
                git_display = html.escape(git_url) if git_url else '未填写'
                git_html = f'<a href="{html.escape(git_url)}" target="_blank">{git_display}</a>' if git_url else git_display
                desc = html.escape(p['description'] or '未填写')

                proj_cards += f'''
                <div class="proj-card">
                    <div class="proj-title">
                        <span class="proj-name">{html.escape(p['name'])}{version}</span>
                        <span class="tag {status_class}">{html.escape(status)}</span>
                    </div>
                    <div class="proj-meta">
                        <div class="meta-row"><span class="meta-label">技术栈</span><span class="meta-value">{tech}</span></div>
                        <div class="meta-row"><span class="meta-label">Git地址</span><span class="meta-value">{git_html}</span></div>
                        <div class="meta-row"><span class="meta-label">简介</span><span class="meta-value">{desc}</span></div>
                    </div>
                </div>'''

            # 递归子模块
            subs = conn.execute("SELECT * FROM modules WHERE parent_id=?", (module_id,)).fetchall()
            sub_html = ''
            for sub in subs:
                sub_html += build_module_html(sub['id'], depth + 1)

            return f'''
            <div class="module-section" style="margin-left:{indent}px">
                <div class="module-header" onclick="toggleModule(this)">
                    <span class="module-arrow">▼</span>
                    <span class="module-icon">📂</span>
                    <span class="module-name">{safe_name}</span>
                    <span class="module-count">{count_descendant_projects(conn, module_id)}个项目</span>
                </div>
                <div class="module-body">
                    {proj_cards}
                    {sub_html}
                </div>
            </div>'''

        # 统计项目总数
        total = conn.execute(
            "SELECT COUNT(*) as c FROM projects WHERE module_id IN (SELECT id FROM modules WHERE company_id=?)",
            (cid,)
        ).fetchone()['c']

        # 构建模块树
        top_modules = conn.execute(
            "SELECT * FROM modules WHERE company_id=? AND parent_id IS NULL ORDER BY sort_order, id", (cid,)
        ).fetchall()
        modules_html = ''
        for mod in top_modules:
            modules_html += build_module_html(mod['id'])

        conn.close()

        return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>项目清单 · {html.escape(company_name)}</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:"Segoe UI","Microsoft YaHei",sans-serif;background:#e8ecf3;color:#2c3143;font-size:14px;line-height:1.6}}
.container{{max-width:900px;margin:0 auto;padding:24px 20px}}

/* 顶部 */
.header{{background:linear-gradient(135deg,#4f6ef4,#7b93f7);border-radius:14px;padding:28px 32px;color:#fff;margin-bottom:24px}}
.header h1{{font-size:22px;margin-bottom:6px}}
.header .subtitle{{font-size:13px;opacity:.85}}
.header .meta{{font-size:12px;opacity:.65;margin-top:10px}}

/* 模块 */
.module-section{{margin-bottom:8px}}
.module-header{{display:flex;align-items:center;gap:8px;padding:10px 14px;background:#fff;border-radius:10px;cursor:pointer;font-size:14px;font-weight:600;box-shadow:0 1px 4px rgba(0,0,0,.04);transition:.15s;user-select:none}}
.module-header:hover{{background:#f5f7ff}}
.module-arrow{{font-size:11px;color:#8c93a8;transition:transform .2s}}
.module-header.collapsed .module-arrow{{transform:rotate(-90deg)}}
.module-icon{{font-size:16px}}
.module-name{{flex:1}}
.module-count{{font-size:11px;color:#8c93a8;font-weight:400}}
.module-body{{padding:8px 0 8px 12px}}
.module-header.collapsed + .module-body{{display:none}}

/* 项目卡片 */
.proj-card{{background:#fff;border-radius:10px;padding:16px 18px;margin-bottom:8px;box-shadow:0 1px 4px rgba(0,0,0,.04);border:1px solid #f0f2f5}}
.proj-title{{display:flex;align-items:center;gap:10px;margin-bottom:10px}}
.proj-name{{font-size:15px;font-weight:600}}

/* 标签 */
.tag{{display:inline-block;padding:2px 8px;border-radius:4px;font-size:11px;font-weight:500}}
.tag-green{{background:#eef9e6;color:#67c23a}}
.tag-orange{{background:#fef6e6;color:#e6a23c}}
.tag-gray{{background:#f4f4f5;color:#909399}}

/* 元信息 */
.proj-meta{{display:flex;flex-direction:column;gap:6px}}
.meta-row{{display:flex;gap:8px;font-size:13px}}
.meta-label{{color:#8c93a8;min-width:56px;flex-shrink:0}}
.meta-value{{color:#3c4257;word-break:break-all}}
.meta-value a{{color:#4f6ef4;text-decoration:none}}
.meta-value a:hover{{text-decoration:underline}}

/* 底部 */
.footer{{text-align:center;font-size:12px;color:#8c93a8;margin-top:32px;padding:16px 0}}

/* 响应式 */
@media(max-width:640px){{
    .container{{padding:16px 12px}}
    .header{{padding:20px}}
    .header h1{{font-size:18px}}
    .proj-card{{padding:12px 14px}}
}}
</style>
</head>
<body>
<div class="container">
    <div class="header">
        <h1>项目清单</h1>
        <div class="subtitle">{html.escape(company_name)}{(" · " + html.escape(department)) if department else ""}</div>
        <div class="meta">生成时间：{now}</div>
    </div>
    {modules_html}
    <div class="footer">共 {total} 个项目 · 生成于 {now}</div>
</div>
<script>
function toggleModule(el){{ el.classList.toggle('collapsed'); }}
</script>
</body>
</html>'''
